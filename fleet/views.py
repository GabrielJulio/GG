from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import logout
from django.contrib.auth.decorators import login_required
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from .models import Viagem, Motorista, Veiculo
from .form import VeiculoForm, MotoristaForm, ViagemForm


from docx import Document


def home(request):
    # Ordenadas as viagens por data, da mais recente para a mais antiga
    viagens = Viagem.objects.all().order_by('-data')
    veiculos = Veiculo.objects.all().order_by('nome')
    motoristas = Motorista.objects.all().order_by('nome')
    termo_busca = request.POST.get('pesquisa')
    page = request.GET.get('page', 1)
    if termo_busca != None:
        viagens = viagens.filter(motorista__nome__icontains=termo_busca) | viagens.filter(veiculo__nome__icontains=termo_busca) | viagens.filter(veiculo__marca__icontains=termo_busca) | viagens.filter(veiculo__tipo_veiculo__icontains=termo_busca) | viagens.filter(observacoes__icontains=termo_busca)
        veiculos = veiculos.filter(nome__icontains=termo_busca) | veiculos.filter(tipo_veiculo__icontains=termo_busca) | veiculos.filter(marca__icontains=termo_busca) | veiculos.filter(ano__icontains=termo_busca) | veiculos.filter(placa__icontains=termo_busca) | veiculos.filter(observacoes__icontains=termo_busca)
        motoristas = motoristas.filter(nome__icontains=termo_busca) | motoristas.filter(observacoes__icontains=termo_busca)

    # Filtro por data sempre sera aplicado independente da busca
    data_inicial = request.POST.get('data_inicial')
    data_final = request.POST.get('data_final')
    if data_inicial and data_final != '':
        viagens = viagens.filter(data__range=[data_inicial, data_final])
    elif data_inicial != '' and data_final == '':
        viagens = viagens.filter(data__gte=data_inicial)
    elif data_inicial == '' and data_final != '':
        viagens = viagens.filter(data__lte=data_final)

    # Paginador
    paginador_viagem = Paginator(viagens, 4)
    paginador_veiculo = Paginator(veiculos, 3)
    paginador_motorista = Paginator(motoristas, 3)

    try:
        viagens = paginador_viagem.page(page)
    except PageNotAnInteger:
        viagens = paginador_viagem.page(1)
    except EmptyPage:
        viagens = Viagem.objects.filter(id__contains='s1fs3a51f53as1f5s1a5f31')

    try:
        veiculos = paginador_veiculo.page(page)
    except PageNotAnInteger:
        veiculos = paginador_veiculo.page(1)
    except EmptyPage:
        veiculos = Veiculo.objects.filter(nome__contains='s1fs3a51f53as1f5s1a5f31')

    try:
        motoristas = paginador_motorista.page(page)
    except PageNotAnInteger:
        motoristas = paginador_motorista.page(1)
    except EmptyPage:
        motoristas = Motorista.objects.filter(nome__contains='s1fs3a51f53as1f5s1a5f31')

    # Criando relatorio das buscas em docx
    document = Document()
    viagens = Viagem.objects.all().order_by('data')
    for motorista in motoristas:
        document.add_heading(f'{motorista}', 0)
        resultados_tabela = []
        for viagem in viagens:
            if viagem.motorista == motorista:
                resultados_tabela.append((
                    viagem.data.strftime("%d/%m/%Y"),
                    f"{viagem.saida.strftime('%H:%M')} ~ {viagem.retorno.strftime('%H:%M')}",
                    viagem.observacoes
                ))
        tabela = document.add_table(rows=1, cols=3)
        hdr_cells = tabela.rows[0].cells
        hdr_cells[0].text = 'Data'
        hdr_cells[1].text = 'Horario'
        hdr_cells[2].text = 'Observações'
        for data, horario, observacoes in resultados_tabela:
            row_cells = tabela.add_row().cells
            row_cells[0].text = str(data)
            row_cells[1].text = horario
            row_cells[2].text = observacoes
        document.add_page_break()
    document.add_page_break()
    document.save('staticfiles/docx/relatorio.docx')
    document.save('static/docx/relatorio.docx')
    viagens = Viagem.objects.all().order_by('-data')


    return render(request, 'fleet/index.html', {
        'viagens': viagens,
        'veiculos': veiculos,
        'motoristas': motoristas,
        'termo_busca': termo_busca
    })

def meu_logout(request):
    logout(request)
    return redirect('/')

def teste(request):
    viagens = Viagem.objects.all().order_by('data')
    veiculos = Veiculo.objects.all().order_by('nome')
    motoristas = Motorista.objects.all().order_by('nome')
    termo_busca = request.POST.get('pesquisa')
    page = request.GET.get('page', 1)

    if termo_busca != None:
        viagens = viagens.filter(motorista__nome__icontains=termo_busca) | viagens.filter(veiculo__nome__icontains=termo_busca) | viagens.filter(veiculo__marca__icontains=termo_busca) | viagens.filter(veiculo__tipo_veiculo__icontains=termo_busca) | viagens.filter(observacoes__icontains=termo_busca)
        veiculos = veiculos.filter(nome__icontains=termo_busca) | veiculos.filter(tipo_veiculo__icontains=termo_busca) | veiculos.filter(marca__icontains=termo_busca) | veiculos.filter(ano__icontains=termo_busca) | veiculos.filter(placa__icontains=termo_busca) | veiculos.filter(observacoes__icontains=termo_busca)
        motoristas = motoristas.filter(nome__icontains=termo_busca) | motoristas.filter(observacoes__icontains=termo_busca)

    # Filtro por data sempre sera aplicado independente da busca
    data_inicial = request.POST.get('data_inicial')
    data_final = request.POST.get('data_final')
    if data_inicial and data_final != '':
        viagens = viagens.filter(data__range=[data_inicial, data_final])
    elif data_inicial != '' and data_final == '':
        viagens = viagens.filter(data__gte=data_inicial)
    elif data_inicial == '' and data_final != '':
        viagens = viagens.filter(data__lte=data_final)

    # Paginador
    paginador_viagem = Paginator(viagens, 4)
    paginador_veiculo = Paginator(veiculos, 3)
    paginador_motorista = Paginator(motoristas, 3)

    try:
        viagens = paginador_viagem.page(page)
    except PageNotAnInteger:
        viagens = paginador_viagem.page(1)
    except EmptyPage:
        viagens = Viagem.objects.filter(id__contains='s1fs3a51f53as1f5s1a5f31')

    try:
        veiculos = paginador_veiculo.page(page)
    except PageNotAnInteger:
        veiculos = paginador_veiculo.page(1)
    except EmptyPage:
        veiculos = Veiculo.objects.filter(nome__contains='s1fs3a51f53as1f5s1a5f31')

    try:
        motoristas = paginador_motorista.page(page)
    except PageNotAnInteger:
        motoristas = paginador_motorista.page(1)
    except EmptyPage:
        motoristas = Motorista.objects.filter(nome__contains='s1fs3a51f53as1f5s1a5f31')

    # Criando relatorio das buscas em docx
    document = Document()
    for motorista in motoristas:
        document.add_heading(f'{motorista}', 0)
        resultados_tabela = []
        for viagem in viagens:
            if viagem.motorista == motorista:
                resultados_tabela.append((
                    viagem.data.strftime("%d/%m/%Y"),
                    f"{viagem.saida.strftime('%H:%M')} ~ {viagem.retorno.strftime('%H:%M')}",
                    viagem.observacoes
                    ))
        tabela = document.add_table(rows=1, cols=3)
        hdr_cells = tabela.rows[0].cells
        hdr_cells[0].text = 'Data'
        hdr_cells[1].text = 'Horario'
        hdr_cells[2].text = 'Observações'
        for data, horario, observacoes in resultados_tabela:
            row_cells = tabela.add_row().cells
            row_cells[0].text = str(data)
            row_cells[1].text = horario
            row_cells[2].text = observacoes
        document.add_page_break()
    document.add_page_break()
    document.save('staticfiles/docx/realatorio.docx')
    document.save('static/docx/realatorio.docx')
    x = 1 + 1
    return render(request, "teste.html", {
        'viagens': viagens,
        'veiculos': veiculos,
        'motoristas': motoristas,
        'termo_busca': termo_busca,
    })

@login_required
def nova_viagem(request):
    form = ViagemForm(request.POST or None)
    if form.is_valid():
        form.save()
        return redirect('home')
    return render(request, 'fleet/formularios/form_cadastro.html', {'form': form})


@login_required
def atualizar_viagem(request, id):
    viagem = get_object_or_404(Viagem, pk=id)
    form = ViagemForm(request.POST or None, instance=viagem)
    if form.is_valid():
        form.save()
        return redirect('home')
    return render(request, 'fleet/formularios/form_cadastro.html', {'form': form})


@login_required
def apagar_viagem(request, id):
    viagem = Viagem.objects.get(pk=id)
    if request.method == 'POST':
        viagem.delete()
        return redirect('home')
    return render(request, 'fleet/delete/confirmacao.html', {'viagem': viagem})


@login_required
def novo_motorista(request):
    form = MotoristaForm(request.POST or None)
    if form.is_valid():
        form.save()
        return redirect('home')
    return render(request, 'fleet/formularios/form_cadastro.html', {'form': form})


@login_required
def atualizar_motorista(request, id):
    motorista = get_object_or_404(Motorista, pk=id)
    form = MotoristaForm(request.POST or None, instance=motorista)
    if form.is_valid():
        form.save()
        return redirect('home')
    return render(request, 'fleet/formularios/form_cadastro.html', {'form': form})


@login_required
def apagar_motorista(request, id):
    motorista = Motorista.objects.get(pk=id)
    if request.method == "POST":
        motorista.delete()
        return redirect('home')
    return render(request, 'fleet/delete/confirmacao.html', {'motorista': motorista})


@login_required
def novo_veiculo(request):
    form = VeiculoForm(request.POST or None)
    if form.is_valid():
        form.save()
        return redirect('home')
    return render(request, 'fleet/formularios/form_cadastro.html', {'form': form})


@login_required
def atualizar_veiculo(request, id):
    veiculo = get_object_or_404(Veiculo, pk=id)
    form = VeiculoForm(request.POST or None, instance=veiculo)
    if form.is_valid():
        form.save()
        return redirect('home')
    return render(request, 'fleet/formularios/form_cadastro.html', {'form': form})


@login_required
def apagar_veiculo(request, id):
    veiculo = Veiculo.objects.get(pk=id)
    if request.method == 'POST':
        veiculo.delete()
        return redirect('home')
    return render(request, 'fleet/delete/confirmacao.html', {'veiculo': veiculo})
